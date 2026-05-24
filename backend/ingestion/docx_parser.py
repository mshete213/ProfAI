import io

from docx import Document as DocxDocument

from core.chunker import Chunk, chunk_by_paragraphs


def _extract_tables(doc: DocxDocument) -> list[str]:
    blocks = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n".join(rows))
    return blocks


def chunk_docx(data: bytes, filename: str, title: str | None = None) -> list[Chunk]:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = _extract_tables(doc)
    full_text = "\n\n".join(paragraphs + tables)

    return chunk_by_paragraphs(
        full_text,
        extra_metadata={
            "filename": filename,
            "title": title or filename,
            "source_type": "docx",
        },
    )
